import re
import os
import argparse
from docx import Document
from odf.opendocument import load, OpenDocumentText
from odf.text import P
from odf.table import Table, TableRow, TableCell
from odf.element import Text


def replace_placeholder_in_doc(input_string, input_doc_path):
    # Generate output document path
    base, ext = os.path.splitext(input_doc_path)
    output_doc_path = f"{base}_modified{ext}"

    # Determine file type by extension
    if ext == '.docx':
        replace_in_docx(input_string, input_doc_path, output_doc_path)
    elif ext == '.odt':
        replace_in_odt(input_string, input_doc_path, output_doc_path)
    else:
        raise ValueError("Unsupported file type. Please use a .docx or .odt file.")

    print(f"Output saved to {output_doc_path}")


def replace_in_docx(input_string, input_doc_path, output_doc_path):
    # Load the document
    doc = Document(input_doc_path)

    # Define the regex pattern to find text between two '@' characters
    pattern = re.compile(r'@([^@]*)@')

    # Iterate through each paragraph in the document
    for para in doc.paragraphs:
        if pattern.search(para.text):
            # Replace the matched text with the input string
            para.text = pattern.sub(input_string, para.text)

    # Iterate through each table in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if pattern.search(para.text):
                        para.text = pattern.sub(input_string, para.text)

    # Save the modified document
    doc.save(output_doc_path)


def replace_in_odt(input_string, input_doc_path, output_doc_path):
    # Load the document
    doc = load(input_doc_path)

    # Define the regex pattern to find text between two '@' characters
    pattern = re.compile(r'@([^@]*)@')

    # Create a new OpenDocumentText to store the modified content
    new_doc = OpenDocumentText()

    # Iterate through each paragraph in the document
    for paragraph in doc.getElementsByType(P):
        new_para = P()
        for text in paragraph.childNodes:
            if text.nodeType == Text.TEXT_NODE and pattern.search(text.data):
                # Replace the matched text with the input string
                text.data = pattern.sub(input_string, text.data)
            new_para.addElement(Text(text.data))
        new_doc.text.addElement(new_para)

    # Iterate through each table in the document
    for table in doc.getElementsByType(Table):
        new_table = Table()
        for row in table.getElementsByType(TableRow):
            new_row = TableRow()
            for cell in row.getElementsByType(TableCell):
                new_cell = TableCell()
                for paragraph in cell.getElementsByType(P):
                    new_para = P()
                    for text in paragraph.childNodes:
                        if text.nodeType == Text.TEXT_NODE and pattern.search(text.data):
                            text.data = pattern.sub(input_string, text.data)
                        new_para.addElement(Text(text.data))
                    new_cell.addElement(new_para)
                new_row.addElement(new_cell)
            new_table.addElement(new_row)
        new_doc.spreadsheet.addElement(new_table)

    # Save the modified document
    new_doc.save(output_doc_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Replace placeholders in a document.")
    parser.add_argument("input_string", type=str, help="The string to replace the placeholders with.")
    parser.add_argument("input_doc_path", type=str, help="Path to the input document (either .docx or .odt).")

    args = parser.parse_args()

    replace_placeholder_in_doc(args.input_string, args.input_doc_path)
